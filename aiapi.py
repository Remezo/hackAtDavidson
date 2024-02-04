
from openai import OpenAI
from dotenv import load_dotenv
import os
import config

# load_dotenv()

# api_key = os.getenv("OPENAI_API_KEY")
api_key= config.DevelopmentConfig.OPENAI_API_KEY

client = OpenAI(api_key=api_key)
from docx import Document

# def transcribe_audio(audio_file_path):
#     with open(audio_file_path, 'rb') as audio_file:
#         audio_content = audio_file.read()  # Read the content of the audio file
        
#         transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file, response_format="text")
#     return transcription

# # transcript = transcribe_audio('song.mp3')
# # print(transcript)


def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)

    minutes= {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items
        
    }
    # doc= save_as_docx(minutes, 'meeting_minutes.docx')

    # doc.save("minutes.docx")
    
    return minutes

def abstract_summary_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    summary_content = response.choices[0].message.content
    return summary_content


def key_points_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    key_points = response.choices[0].message.content
    return key_points


def action_item_extraction(transcription):
    response = client.chat.completions.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done if there is any. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    action_items = response.choices[0].message.content
    return action_items

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    
    return doc


# audio_file_path = "song.mp3"
# transcription = transcribe_audio(audio_file_path)

# with open("transcript.txt", "r") as file:
#     transcription = file.read()

# minutes = meeting_minutes(transcription)
# print(minutes)
# save_as_docx(minutes, 'meeting_minutes.docx')
